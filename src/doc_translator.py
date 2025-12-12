# src/doc_translator.py

import io
import docx
import fitz  # PyMuPDF
import pandas as pd
import zipfile
import tempfile
import os
from fpdf import FPDF
import xlwings as xw
import duckdb  
from transformers import pipeline

from security_config import TEMP_DIR

# ==========================================
# PARTIE 1 — Dictionnaire des modèles
# ==========================================
MODEL_MAP = {
    "Deutsch": "Helsinki-NLP/opus-mt-fr-de",
    "English": "Helsinki-NLP/opus-mt-fr-en",
    "Español": "Helsinki-NLP/opus-mt-fr-es",
    "Português": "Helsinki-NLP/opus-mt-fr-pt",
    "Nederlands": "Helsinki-NLP/opus-mt-fr-nl",
    "Русский": "Helsinki-NLP/opus-mt-fr-ru",
    "Polski": "Helsinki-NLP/opus-mt-fr-pl",
    "Svenska": "Helsinki-NLP/opus-mt-fr-sv",
    "Dansk": "Helsinki-NLP/opus-mt-fr-da",
    "Norsk": "Helsinki-NLP/opus-mt-fr-no",
    "Suomi": "Helsinki-NLP/opus-mt-fr-fi",
    "Čeština": "Helsinki-NLP/opus-mt-fr-cs",
    "Slovenčina": "Helsinki-NLP/opus-mt-fr-sk",
    "Magyar": "Helsinki-NLP/opus-mt-fr-hu",
    "Română": "Helsinki-NLP/opus-mt-fr-ro",
    "Български": "Helsinki-NLP/opus-mt-fr-bg",
    "Українська": "Helsinki-NLP/opus-mt-fr-uk",
    "Ελληνικά": "Helsinki-NLP/opus-mt-fr-el",
    "Türkçe": "Helsinki-NLP/opus-mt-fr-tr",
    "Italiano": "Helsinki-NLP/opus-mt-fr-it",
    "中文": "Helsinki-NLP/opus-mt-fr-zh",
    "日本語": "Helsinki-NLP/opus-mt-fr-ja",
    "한국어": "Helsinki-NLP/opus-mt-fr-ko",
    "हिन्दी": "Helsinki-NLP/opus-mt-fr-hi",
    "العربية": "Helsinki-NLP/opus-mt-fr-ar",
    "فارسی": "Helsinki-NLP/opus-mt-fr-fa",
    "Bahasa Indonesia": "Helsinki-NLP/opus-mt-fr-id",
    "Bahasa Melayu": "Helsinki-NLP/opus-mt-fr-ms",
    "Kiswahili": "Helsinki-NLP/opus-mt-fr-sw",
    "Yorùbá": "Helsinki-NLP/opus-mt-fr-yo",
    "Igbo": "Helsinki-NLP/opus-mt-fr-ig",
    "Hausa": "Helsinki-NLP/opus-mt-fr-ha",
    "isiZulu": "Helsinki-NLP/opus-mt-fr-zu",
    "አማርኛ": "Helsinki-NLP/opus-mt-fr-am",
    "Kreyòl Ayisyen": "Helsinki-NLP/opus-mt-fr-ht",
    "Runa Simi": "Helsinki-NLP/opus-mt-fr-qu",
    "Avañe'ẽ": "Helsinki-NLP/opus-mt-fr-gn",
}

# ==========================================
# PARTIE 2 — Fonctions de traduction (Base, PDF, DOCX, TXT, ZIP)
# ==========================================

def translate_text(text: str, translator, chunk_size=1000) -> str:
    """Traduit le texte par morceaux."""
    if not text:
        return ""
    text = str(text)
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    out = []
    for ch in chunks:
        # Assurez-vous que le modèle est bien chargé
        if isinstance(translator, pipeline):
            res = translator(ch, max_length=4000)[0]["translation_text"]
            out.append(res)
        else:
            # Si le modèle n'est pas chargé (pour les tests), renvoyer le texte original
            out.append(ch) 
    return "\n".join(out)

# ... (Insérer les fonctions translate_pdf_bytes, translate_docx_bytes, translate_txt_bytes) ...
# Les fonctions ci-dessous utilisent l'implémentation de ton `translator.py` original :

def translate_pdf_bytes(pdf_bytes: bytes, translator):
    # ... (Ton code pour PDF) ...
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = "".join([page.get_text() for page in doc])
    translated = translate_text(text, translator)

    pdf_out = FPDF()
    pdf_out.add_page()
    pdf_out.set_font("Arial", size=12)
    safe_text = translated.encode('latin-1', 'replace').decode('latin-1')
    for line in safe_text.split("\n"):
        pdf_out.multi_cell(0, 8, line)
    buf = io.BytesIO()
    pdf_out.output(buf)
    buf.seek(0)
    return buf, "document_traduit.pdf", "application/pdf"

def translate_docx_bytes(docx_bytes: bytes, translator):
    # ... (Ton code pour DOCX) ...
    in_buf = io.BytesIO(docx_bytes)
    doc_in = docx.Document(in_buf)
    for p in doc_in.paragraphs:
        if p.text.strip():
            p.text = translate_text(p.text, translator)
    # ... (Traitement des tables) ...
    for table in doc_in.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = translate_text(cell.text, translator)
    out_buf = io.BytesIO()
    doc_in.save(out_buf)
    out_buf.seek(0)
    return out_buf, "document_traduit.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

def translate_txt_bytes(txt_bytes: bytes, translator):
    # ... (Ton code pour TXT) ...
    try:
        text = txt_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = txt_bytes.decode("latin-1", errors="ignore")
    translated = translate_text(text, translator)
    out_buf = io.BytesIO(translated.encode("utf-8"))
    return out_buf, "document_traduit.txt", "text/plain"

# ==========================================
# PARTIE 3 — Optimisation DuckDB & xlwings (CSV/EXCEL)
# ==========================================

def translate_excel_with_app(x_bytes: bytes, translator, original_filename: str):
    """Gère .xls et .xlsx en utilisant xlwings et le dossier temporaire sécurisé."""
    # ... (Ton code Excel) ...
    lower = original_filename.lower()
    if lower.endswith(".xlsx"):
        suffix = ".xlsx"
        mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        suffix = ".xls"
        mime_type = "application/vnd.ms-excel"

    # Utilisation du chemin sécurisé (TEMP_DIR)
    tmp_in = tempfile.NamedTemporaryFile(delete=False, suffix=suffix, dir=TEMP_DIR) 
    tmp_in.write(x_bytes)
    tmp_in.close() 

    input_path = os.path.abspath(tmp_in.name)

    app = xw.App(visible=False)
    app.display_alerts = False
    
    try:
        # ... (Le reste de ton code xlwings) ...
        wb_in = app.books.open(input_path)
        wb_out = app.books.add()

        for sheet in wb_in.sheets:
            sheet.api.Copy(Before=wb_out.sheets[0].api)
            target_sheet = wb_out.sheets[0]
            
            try:
                target_sheet.api.Unprotect()
            except Exception:
                pass

            try:
                used_range = target_sheet.used_range
                for cell in used_range:
                    try:
                        if cell.value and isinstance(cell.value, str) and cell.value.strip():
                            cell.value = translate_text(cell.value, translator)
                    except Exception:
                        pass
            except Exception:
                pass

        if len(wb_out.sheets) > 1:
            try:
                wb_out.sheets[-1].delete()
            except:
                pass

        tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix=suffix, dir=TEMP_DIR) # Utilisation de TEMP_DIR
        tmp_out.close()
        output_path = os.path.abspath(tmp_out.name)

        wb_out.save(output_path)
        
        wb_in.close()
        wb_out.close()
        app.quit()

        with open(output_path, "rb") as f:
            result = f.read()

        try:
            os.unlink(input_path)
            os.unlink(output_path)
        except:
            pass
        
    except Exception as e:
        # ... (Gestion des erreurs et nettoyage) ...
        try:
            app.quit()
        except:
            pass
        try:
            os.unlink(input_path)
        except:
            pass
        raise e

    return io.BytesIO(result), f"document_traduit{suffix}", mime_type


def translate_csv_bytes_duckdb(csv_bytes: bytes, translator):
    """Utilise DuckDB pour une lecture/écriture ultra-rapide du CSV, utilisant le chemin sécurisé."""
    
    # 1. Écriture temporaire du fichier source dans le dossier sécurisé
    tmp_in = tempfile.NamedTemporaryFile(delete=False, suffix=".csv", dir=TEMP_DIR) # Utilisation de TEMP_DIR
    tmp_in.write(csv_bytes)
    tmp_in.close()
    input_path = os.path.abspath(tmp_in.name)
    
    try:
        # 2. Lecture optimisée avec DuckDB
        query = f"SELECT * FROM read_csv_auto('{input_path}', sep=';', all_varchar=True)"
        df = duckdb.query(query).to_df()
        
        # 3. Traduction (opération lourde en Python)
        for col in df.columns:
            # Correction: on utilise ton séparateur ";" pour les colonnes Excel
            if not pd.api.types.is_numeric_dtype(df[col]):
                df[col] = df[col].astype(str).apply(lambda x: translate_text(x, translator))
            
        # 4. Écriture optimisée via DuckDB (COPY TO)
        tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix=".csv", dir=TEMP_DIR) # Utilisation de TEMP_DIR
        tmp_out.close()
        output_path = os.path.abspath(tmp_out.name)
        
        duckdb.register("df_translated", df)
        # Utilise ton séparateur Excel/CSV personnel : ";"
        duckdb.sql(f"COPY df_translated TO '{output_path}' (HEADER, DELIMITER ';')") 
        
        # Lecture du résultat
        with open(output_path, "rb") as f:
            result = f.read()
            
        # Nettoyage
        os.unlink(input_path)
        os.unlink(output_path)
        
        return io.BytesIO(result), "document_traduit.csv", "text/csv"

    except Exception as e:
        # ... (Gestion des erreurs et nettoyage) ...
        try:
            os.unlink(input_path)
        except:
            pass
        raise e

# -------------------------
# Dispatcher pour Excel/CSV
# -------------------------
def dispatch_excel_csv(file_bytes: bytes, filename: str, translator):
    """Dirige vers la fonction CSV (DuckDB) ou Excel (xlwings)."""
    lower = filename.lower()
    if lower.endswith(".csv"):
        return translate_csv_bytes_duckdb(file_bytes, translator) 
    elif lower.endswith(".xlsx") or lower.endswith(".xls"):
        return translate_excel_with_app(file_bytes, translator, filename)
    raise ValueError(f"Format de fichier de données non reconnu: {filename}")

# -------------------------
# ZIP (qui utilise le dispatcher)
# -------------------------
def translate_zip_bytes(zip_bytes: bytes, translator):
    # ... (Ton code ZIP, utilisant dispatch_excel_csv) ...
    in_buf = io.BytesIO(zip_bytes)
    out_buf = io.BytesIO()
    with zipfile.ZipFile(in_buf, "r") as zin, zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name in zin.namelist():
            with zin.open(name) as f:
                content = f.read()
            lower = name.lower()
            try:
                if lower.endswith(".pdf"):
                    translated_file, _, _ = translate_pdf_bytes(content, translator)
                    zout.writestr(name, translated_file.getvalue())
                elif lower.endswith(".docx"):
                    translated_file, _, _ = translate_docx_bytes(content, translator)
                    zout.writestr(name, translated_file.getvalue())
                elif lower.endswith(".txt"):
                    translated_file, _, _ = translate_txt_bytes(content, translator)
                    zout.writestr(name, translated_file.getvalue())
                elif lower.endswith(".xlsx") or lower.endswith(".xls") or lower.endswith(".csv"):
                    # Utilise le dispatcher pour traiter les formats de données
                    translated_file, _, _ = dispatch_excel_csv(content, name, translator) 
                    zout.writestr(name, translated_file.getvalue())
                else:
                    zout.writestr(name, content)
            except Exception:
                zout.writestr(name, content)
    out_buf.seek(0)
    return out_buf, "documents_traduits.zip", "application/zip"


# -------------------------
# Dispatcher Principal (Exposé à Streamlit)
# -------------------------
def dispatch_file_for_translation(file_bytes: bytes, filename: str, translator):
    """Fonction principale appelée par l'interface Streamlit."""
    filename_lower = filename.lower()
    
    if filename_lower.endswith(".pdf"):
        return translate_pdf_bytes(file_bytes, translator)
    elif filename_lower.endswith(".docx"):
        return translate_docx_bytes(file_bytes, translator)
    elif filename_lower.endswith(".txt"):
        return translate_txt_bytes(file_bytes, translator)
    elif filename_lower.endswith(".xlsx") or filename_lower.endswith(".xls") or filename_lower.endswith(".csv"):
        # Utilise le dispatcher Excel/CSV
        return dispatch_excel_csv(file_bytes, filename, translator) 
    elif filename_lower.endswith(".zip"):
        return translate_zip_bytes(file_bytes, translator)
    else:
        raise ValueError("Format non supporté.")